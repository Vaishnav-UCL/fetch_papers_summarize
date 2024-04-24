import tkinter as tk
from tkinter import simpledialog, messagebox
import pandas as pd
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
import time
import re
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize, sent_tokenize
from nltk.probability import FreqDist
from docx import Document

nltk.download('punkt')
nltk.download('stopwords')

def summarize_text(text):
    stop_words = set(stopwords.words('english'))
    words = word_tokenize(text.lower())
    freq_dist = FreqDist(word for word in words if word not in stop_words and word.isalnum())
    important_sentences = {}
    sents = sent_tokenize(text)
    for i, sent in enumerate(sents):
        for word in word_tokenize(sent.lower()):
            if word in freq_dist:
                if i in important_sentences:
                    important_sentences[i] += freq_dist[word]
                else:
                    important_sentences[i] = freq_dist[word]
    indexes = sorted(important_sentences, key=important_sentences.get, reverse=True)[:3]
    summary = ' '.join([sents[idx] for idx in sorted(indexes)])
    return summary

def fetch_papers(scientist_name, years, keywords):
    results = []
    options = Options()
    driver = webdriver.Chrome(options=options)
    current_year = pd.Timestamp.now().year

    try:
        driver.get(f"https://scholar.google.com/scholar?q={scientist_name}")
        WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH, '//*[@id="gs_res_ccl_mid"]/div[1]/table/tbody/tr/td[2]/h4/a'))).click()
        WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.LINK_TEXT, 'YEAR'))).click()
        time.sleep(5)

        articles = WebDriverWait(driver, 20).until(
            EC.presence_of_all_elements_located((By.XPATH, '//*[@id="gsc_a_b"]/tr/td[1]/a')))

        for article in articles[:4]:  # Limiting to first 4 articles
            driver.execute_script("window.open(arguments[0]);", article.get_attribute('href'))
            driver.switch_to.window(driver.window_handles[1])

            try:
                WebDriverWait(driver, 10).until(
                    EC.presence_of_element_located((By.XPATH, '//*[@id="gsc_oci_descr"]')))
                abstract = driver.find_element(By.XPATH, '//*[@id="gsc_oci_descr"]').text
                title = driver.find_element(By.XPATH, '//*[@id="gsc_oci_title"]/a').text
                pub_info = driver.find_element(By.XPATH, '//*[@id="gsc_oci_table"]/div[2]/div[2]').text
                publication_year_match = re.search(r'\b(20\d{2})\b', pub_info)

                if publication_year_match:
                    publication_year = int(publication_year_match.group())
                else:
                    print(f"No valid year found for {title}. Skipping...")
                    continue

                if publication_year >= current_year - years:
                    if not keywords or any(keyword in title.lower() or keyword in abstract.lower() for keyword in keywords):
                        summary = summarize_text(abstract)
                        results.append({"Scientist Name": scientist_name, "Title": title, "Year": publication_year, "Abstract Summary": summary})
                    else:
                        print(f"Article '{title}' excluded based on keywords.")
                else:
                    print(f"Article '{title}' excluded based on year.")

            except Exception as e:
                print(f"Error processing article: {str(e)}")
            finally:
                driver.close()
                driver.switch_to.window(driver.window_handles[0])

    finally:
        driver.quit()

    return results

def save_to_excel(results, scientist_name):
    df = pd.DataFrame(results)
    df.to_excel(f"{scientist_name}_papers.xlsx", index=False)

def save_to_word(results, scientist_name):
    doc = Document()
    doc.add_heading(f'Research Summary for {scientist_name}', 0)
    for result in results:
        doc.add_heading(result["Title"], level=1)
        doc.add_paragraph(f"Year: {result['Year']}")
        doc.add_paragraph("Summary:")
        doc.add_paragraph(result["Abstract Summary"])
    doc.save(f"{scientist_name}_summary.docx")

def main():
    root = tk.Tk()
    root.withdraw()

    scientist_name = simpledialog.askstring("Input", "Enter the scientist's name:", parent=root)
    years = simpledialog.askinteger("Input", "Enter the number of years (e.g., 5 for the last 5 years):", parent=root)
    keywords = simpledialog.askstring("Input", "Enter keywords separated by commas (leave blank for all):", parent=root)
    keywords = [keyword.strip().lower() for keyword in keywords.split(',')] if keywords else []

    if scientist_name and years is not None:
        results = fetch_papers(scientist_name, years, keywords)
        if results:
            save_to_excel(results, scientist_name)
            save_to_word(results, scientist_name)
            messagebox.showinfo("Success", "Data written to Excel and Word.")
        else:
            messagebox.showinfo("No Results", "No matching papers found.")
    else:
        messagebox.showerror("Error", "Scientist name and years are required!")

if __name__ == "__main__":
    main()

